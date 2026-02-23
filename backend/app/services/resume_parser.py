from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

ALLOWED_EXTENSIONS = {".pdf", ".docx"}


def _parse_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages).strip()


def _parse_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)

    return "\n".join(lines).strip()


def parse_resume_file(filename: str, content: bytes) -> str:
    extension = Path(filename).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")

    if not content:
        raise ValueError("Uploaded file is empty.")

    if extension == ".pdf":
        text = _parse_pdf(content)
    else:
        text = _parse_docx(content)

    if len(text) < 20:
        raise ValueError(
            "Could not extract enough text from the file. "
            "Try a text-based PDF/DOCX (not scanned image-only)."
        )

    return text
